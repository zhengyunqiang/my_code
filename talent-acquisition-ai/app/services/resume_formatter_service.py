"""
Resume Formatting Service - Phase 2 Module.
Handles resume standardization, privacy protection, and format conversion.
"""
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

import aiofiles
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pypdf import PdfReader
from PIL import Image
import pytesseract

from app.core.config import get_settings
from app.core.logger import get_logger

settings = get_settings()
logger = get_logger(__name__)


class PrivacyProtector:
    """Protect candidate privacy by removing sensitive information."""

    # Patterns for sensitive information
    PRIVACY_PATTERNS = {
        "phone": [
            r'(?:(?:\+|00)86)?1[3-9]\d{9}',
            r'\d{3}-\d{4}-\d{4}',
            r'\d{3}\d{4}\d{4}',
        ],
        "email": [
            r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
        ],
        "wechat": [
            r'(?:微信|WeChat|wechat)[:：\s]*[A-Za-z0-9_-]+',
            r'wxid_[A-Za-z0-9_-]+',
        ],
        "id_card": [
            r'\d{17}[\dXx]',  # Chinese ID card
            r'[A-Z]{2}\d{6}\(\d\)',  # Hong Kong ID
        ],
        "address": [
            r'(?:地址|住址|Address)[:：][^\n]{10,100}',
        ],
    }

    @staticmethod
    def redact_phone(text: str, replacement: str = "[手机号已保护]") -> str:
        """Redact phone numbers from text."""
        for pattern in PrivacyProtector.PRIVACY_PATTERNS["phone"]:
            text = re.sub(pattern, replacement, text)
        return text

    @staticmethod
    def redact_email(text: str, replacement: str = "[邮箱已保护]") -> str:
        """Redact email addresses from text."""
        for pattern in PrivacyProtector.PRIVACY_PATTERNS["email"]:
            text = re.sub(pattern, replacement, text)
        return text

    @staticmethod
    def redact_wechat(text: str, replacement: str = "[微信号已保护]") -> str:
        """Redact WeChat IDs from text."""
        for pattern in PrivacyProtector.PRIVACY_PATTERNS["wechat"]:
            text = re.sub(pattern, replacement, text)
        return text

    @staticmethod
    def redact_id_card(text: str, replacement: str = "[身份证号已保护]") -> str:
        """Redact ID card numbers from text."""
        for pattern in PrivacyProtector.PRIVACY_PATTERNS["id_card"]:
            text = re.sub(pattern, replacement, text)
        return text

    @staticmethod
    def redact_address(text: str, replacement: str = "[地址已保护]") -> str:
        """Redact addresses from text."""
        for pattern in PrivacyProtector.PRIVACY_PATTERNS["address"]:
            text = re.sub(pattern, replacement, text)
        return text

    @classmethod
    def redact_all(cls, text: str) -> tuple[str, dict[str, Any]]:
        """
        Redact all sensitive information from text.

        Args:
            text: Input text

        Returns:
            Tuple of (redacted text, metadata about redacted items)
        """
        metadata = {
            "phones_found": len(re.findall(cls.PRIVACY_PATTERNS["phone"][0], text)),
            "emails_found": len(re.findall(cls.PRIVACY_PATTERNS["email"][0], text)),
            "wechats_found": 0,
            "id_cards_found": 0,
        }

        redacted = text
        redacted = cls.redact_phone(redacted)
        redacted = cls.redact_email(redacted)
        redacted = cls.redact_wechat(redacted)
        redacted = cls.redact_id_card(redacted)
        redacted = cls.redact_address(redacted)

        return redacted, metadata


class ResumeTextExtractor:
    """Extract text from various file formats."""

    @staticmethod
    async def extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            reader = PdfReader(file_path)
            text = ""

            for page in reader.pages:
                text += page.extract_text() + "\n"

            return text.strip()

        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            raise

    @staticmethod
    async def extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = ""

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "

            return text.strip()

        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise

    @staticmethod
    async def extract_from_image(file_path: str) -> str:
        """Extract text from image using OCR."""
        try:
            image = Image.open(file_path)

            # Configure Tesseract
            pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd

            # Extract text with Chinese and English
            text = pytesseract.image_to_string(
                image,
                lang=settings.tesseract_language,
            )

            return text.strip()

        except Exception as e:
            logger.error(f"Error extracting text from image {file_path}: {e}")
            raise

    @staticmethod
    async def extract_from_txt(file_path: str) -> str:
        """Extract text from plain text file."""
        try:
            async with aiofiles.open(file_path, mode='r', encoding='utf-8') as f:
                return await f.read()

        except UnicodeDecodeError:
            # Try with different encoding
            try:
                async with aiofiles.open(file_path, mode='r', encoding='gbk') as f:
                    return await f.read()
            except Exception as e:
                logger.error(f"Error reading text file {file_path}: {e}")
                raise

    async def extract_text(self, file_path: str) -> str:
        """
        Extract text from file based on extension.

        Args:
            file_path: Path to file

        Returns:
            Extracted text content
        """
        file_ext = Path(file_path).suffix.lower()

        extractors = {
            ".pdf": self.extract_from_pdf,
            ".docx": self.extract_from_docx,
            ".doc": self.extract_from_docx,
            ".txt": self.extract_from_txt,
            ".jpg": self.extract_from_image,
            ".jpeg": self.extract_from_image,
            ".png": self.extract_from_image,
        }

        extractor = extractors.get(file_ext)
        if not extractor:
            raise ValueError(f"Unsupported file format: {file_ext}")

        return await extractor(file_path)


class ResumeFormatter:
    """Format and standardize resumes."""

    STANDARD_TEMPLATE = """
================================================================================
                              标准简历格式
================================================================================

基本信息
--------------------------------------------------------------------------------
姓名: {name}
性别: {gender}
年龄: {age}
电话: {phone}
邮箱: {email}
工作年限: {experience_years}年
最高学历: {education_level}

求职意向
--------------------------------------------------------------------------------
期望职位: {target_position}
期望薪资: {expected_salary}
期望地点: {expected_location}
工作性质: {employment_type}

工作经历
--------------------------------------------------------------------------------
{work_history}

项目经验
--------------------------------------------------------------------------------
{project_experience}

教育背景
--------------------------------------------------------------------------------
{education}

专业技能
--------------------------------------------------------------------------------
{skills}

自我评价
--------------------------------------------------------------------------------
{self_evaluation}

================================================================================
                        简历生成时间: {generation_time}
================================================================================
"""

    async def format_to_standard_text(
        self,
        extracted_data: dict[str, Any],
        additional_notes: Optional[str] = None,
    ) -> str:
        """
        Format extracted data into standard text template.

        Args:
            extracted_data: Dictionary with extracted resume data
            additional_notes: Optional additional notes to append

        Returns:
            Formatted resume text
        """
        # Format work history
        work_history = ""
        if extracted_data.get("work_history"):
            for idx, work in enumerate(extracted_data["work_history"], 1):
                work_history += f"""
{work.get('company', 'Unknown')} | {work.get('position', 'Unknown')}
时间: {work.get('duration', 'Unknown')}
职责: {work.get('responsibilities', 'N/A')}
"""

        # Format skills
        skills = ", ".join(extracted_data.get("skills", []))

        # Format into template
        formatted = self.STANDARD_TEMPLATE.format(
            name=extracted_data.get("name", "[姓名待确认]"),
            gender=extracted_data.get("gender", "[性别待确认]"),
            age=extracted_data.get("age", "[年龄待确认]"),
            phone="[联系方式已保护]",
            email="[邮箱已保护]",
            experience_years=extracted_data.get("years_of_experience", 0),
            education_level=extracted_data.get("education_level", "[学历待确认]"),
            target_position=extracted_data.get("target_position", "[待确认]"),
            expected_salary=extracted_data.get("expected_salary", "面议"),
            expected_location=extracted_data.get("expected_location", "待定"),
            employment_type=extracted_data.get("employment_type", "全职"),
            work_history=work_history or "暂无",
            project_experience=extracted_data.get("project_experience", "暂无"),
            education=extracted_data.get("education", "暂无"),
            skills=skills or "暂无",
            self_evaluation=extracted_data.get("self_evaluation", "暂无"),
            generation_time=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        )

        # Add additional notes if provided
        if additional_notes:
            formatted += f"\n\n补充信息:\n{additional_notes}\n"

        return formatted

    async def format_to_word(
        self,
        extracted_data: dict[str, Any],
        output_path: str,
        additional_notes: Optional[str] = None,
    ) -> None:
        """
        Format resume into Word document with standard template.

        Args:
            extracted_data: Dictionary with extracted resume data
            output_path: Output file path
            additional_notes: Optional additional notes
        """
        doc = Document()

        # Title
        title = doc.add_heading("候选人简历", 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Basic Information
        doc.add_heading("基本信息", level=1)

        basic_info_table = doc.add_table(rows=7, cols=2)
        basic_info_table.style = 'Light Grid Accent 1'

        basic_info_data = [
            ("姓名", extracted_data.get("name", "[待确认]")),
            ("性别", extracted_data.get("gender", "[待确认]")),
            ("年龄", str(extracted_data.get("age", "[待确认]"))),
            ("工作年限", f"{extracted_data.get('years_of_experience', 0)}年"),
            ("最高学历", extracted_data.get("education_level", "[待确认]")),
            ("电话", "[联系方式已保护]"),
            ("邮箱", "[邮箱已保护]"),
        ]

        for idx, (key, value) in enumerate(basic_info_data):
            row = basic_info_table.rows[idx]
            row.cells[0].text = key
            row.cells[1].text = str(value)

        # Job Intention
        doc.add_heading("求职意向", level=1)
        doc.add_paragraph(f"期望职位: {extracted_data.get('target_position', '[待确认]')}")
        doc.add_paragraph(f"期望薪资: {extracted_data.get('expected_salary', '面议')}")
        doc.add_paragraph(f"期望地点: {extracted_data.get('expected_location', '待定')}")
        doc.add_paragraph(f"工作性质: {extracted_data.get('employment_type', '全职')}")

        # Work History
        if extracted_data.get("work_history"):
            doc.add_heading("工作经历", level=1)
            for work in extracted_data["work_history"]:
                p = doc.add_paragraph()
                p.add_run(f"{work.get('company', 'Unknown')} | ").bold = True
                p.add_run(f"{work.get('position', 'Unknown')}")

                doc.add_paragraph(f"工作时间: {work.get('duration', 'Unknown')}")
                doc.add_paragraph(f"工作职责:\n{work.get('responsibilities', 'N/A')}")

        # Skills
        doc.add_heading("专业技能", level=1)
        skills = extracted_data.get("skills", [])
        if skills:
            doc.add_paragraph(", ".join(skills))
        else:
            doc.add_paragraph("暂无")

        # Additional Notes
        if additional_notes:
            doc.add_heading("补充信息", level=1)
            doc.add_paragraph(additional_notes)

        # Footer
        doc.add_paragraph(f"\n简历生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", style='Intense Quote')

        # Save document
        doc.save(output_path)

        logger.info(f"Formatted Word document saved to: {output_path}")


class ResumeFormatterService:
    """Main service for resume formatting operations."""

    def __init__(self):
        self.extractor = ResumeTextExtractor()
        self.privacy_protector = PrivacyProtector()
        self.formatter = ResumeFormatter()

    async def process_resume_file(
        self,
        file_path: str,
        output_dir: Optional[str] = None,
        protect_privacy: bool = True,
        format_to_word: bool = False,
        additional_notes: Optional[str] = None,
    ) -> dict[str, Any]:
        """
        Process a resume file: extract, protect privacy, and format.

        Args:
            file_path: Input resume file path
            output_dir: Optional output directory
            protect_privacy: Whether to redact sensitive information
            format_to_word: Whether to format to Word document
            additional_notes: Optional additional notes

        Returns:
            Processing result with paths and metadata
        """
        # Extract text
        raw_text = await self.extractor.extract_text(file_path)

        # Protect privacy if requested
        if protect_privacy:
            redacted_text, privacy_metadata = self.privacy_protector.redact_all(raw_text)
        else:
            redacted_text = raw_text
            privacy_metadata = {}

        # Extract structured data
        from app.services.resume_screening_service import ResumeParser
        parser = ResumeParser()

        extracted_data = {
            "name": parser._extract_name(redacted_text),
            "phone": self.privacy_protector.redact_phone(parser.extract_phone(raw_text) or ""),
            "email": self.privacy_protector.redact_email(parser.extract_email(raw_text) or ""),
            "age": parser.extract_age(redacted_text),
            "gender": None,  # Would need more sophisticated extraction
            "years_of_experience": parser.extract_experience_years(redacted_text),
            "education_level": str(parser.extract_education_level(redacted_text)),
            "skills": parser.extract_skills(redacted_text),
            "work_history": parser.extract_work_history(redacted_text),
        }

        # Generate formatted text
        formatted_text = await self.formatter.format_to_standard_text(
            extracted_data,
            additional_notes,
        )

        # Determine output paths
        input_path = Path(file_path)
        if output_dir:
            output_path = Path(output_dir)
        else:
            output_path = input_path.parent

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_name = input_path.stem

        # Save formatted text
        txt_output_path = output_path / f"{base_name}_formatted_{timestamp}.txt"
        async with aiofiles.open(txt_output_path, mode='w', encoding='utf-8') as f:
            await f.write(formatted_text)

        # Format to Word if requested
        word_output_path = None
        if format_to_word:
            word_output_path = output_path / f"{base_name}_formatted_{timestamp}.docx"
            await self.formatter.format_to_word(
                extracted_data,
                str(word_output_path),
                additional_notes,
            )

        return {
            "original_file": str(file_path),
            "formatted_text_path": str(txt_output_path),
            "formatted_word_path": str(word_output_path) if word_output_path else None,
            "raw_text": raw_text,
            "redacted_text": redacted_text,
            "extracted_data": extracted_data,
            "privacy_metadata": privacy_metadata,
            "processing_time": datetime.utcnow().isoformat(),
        }


# Global service instance
resume_formatter_service = ResumeFormatterService()
